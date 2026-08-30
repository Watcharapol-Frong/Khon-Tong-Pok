"""
ตัวประกอบร่าง: ไฟล์อะไรก็ได้ที่ผู้ใช้แนบมา -> ข้อความ

ลำดับการเลือกวิธี (สำคัญ อย่าสลับ)
----------------------------------
  1. text layer ของ PDF   ฟรี · เร็ว · ตรงเป๊ะ 100%
  2. Typhoon OCR          เสียโควตา · เข้าใจ layout ไทย · แม่นสุดสำหรับภาพ
  3. Tesseract            ฟรี · ต้องลงเอง · พังกับตารางและสองคอลัมน์

เรซูเม่ที่ export จาก Canva/Word ซึ่งเป็นส่วนใหญ่ จบตั้งแต่ข้อ 1
ไม่ต้องยิง API เลยสักครั้ง — ประหยัดโควตาไว้ให้ไฟล์ที่จำเป็นจริง ๆ
ถ้ายิง OCR ทุกไฟล์แบบไม่คิด โควตาจะหมดตั้งแต่ยังไม่ถึงวัน pitch
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass

from ai.ocr.typhoon_ocr import TyphoonOCR, TyphoonOCRError
from ai.skill.pdf import (
    MIN_CHARS_PER_PAGE,
    tesseract_image,
    tesseract_pdf,
    text_layer,
)

log = logging.getLogger(__name__)

IMAGE_EXT = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff")


@dataclass
class DocumentText:
    text: str
    method: str        # text_layer | typhoon_ocr | tesseract | docx | none
    ocr_used: bool
    pages: int
    note: str = ""

    @property
    def ok(self) -> bool:
        return len(self.text.strip()) >= 40


def _clean(text: str) -> str:
    """
    ล้างตัวคั่นบรรทัดแปลก ๆ ที่ PDF แถมมา

    เหตุผลเดียวกับใน train/ingest_resumes.py — U+2028/U+2029 และเพื่อน ๆ
    ทำให้ json.dumps เขียนไฟล์ที่อ่านกลับไม่ได้ ("Unterminated string")
    ที่นี่ข้อความจะถูกยัดลง JSON response และ jsonb ใน Postgres ด้วย
    เลยต้องล้างที่จุดเดียวกันตั้งแต่แรก
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub("[\u000b\u000c\u001c\u001d\u001e\u0085\u2028\u2029]", "\n", text)
    text = re.sub("[ \t\u00a0\u200b]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_document(
    data: bytes,
    filename: str,
    *,
    typhoon: TyphoonOCR | None = None,
    tesseract_cmd: str = "",
) -> DocumentText:
    name = (filename or "").lower()

    if name.endswith(".docx"):
        return _read_docx(data)

    if name.endswith(IMAGE_EXT):
        return _read_image(data, typhoon=typhoon, tesseract_cmd=tesseract_cmd)

    return _read_pdf(data, typhoon=typhoon, tesseract_cmd=tesseract_cmd)


# ---------------------------------------------------------------------------
def _read_pdf(
    data: bytes, *, typhoon: TyphoonOCR | None, tesseract_cmd: str
) -> DocumentText:
    text, avg, n_pages = text_layer(data)

    if avg >= MIN_CHARS_PER_PAGE:
        return DocumentText(_clean(text), "text_layer", False, n_pages)

    log.info("text layer บางมาก (%.0f ตัว/หน้า) — ต้อง OCR", avg)

    if typhoon and typhoon.available:
        try:
            # ส่ง text layer ที่แกะได้ไปเป็น hint ด้วย ถึงจะน้อยแต่ช่วยเรื่องสะกด
            ocr_text, pages_done = typhoon.ocr_pdf(data, hint=text)
            if ocr_text.strip():
                note = ""
                if n_pages > pages_done:
                    note = f"ไฟล์มี {n_pages} หน้า อ่านให้ {pages_done} หน้าแรก"
                return DocumentText(_clean(ocr_text), "typhoon_ocr", True, pages_done, note)
            log.warning("Typhoon OCR อ่านได้ข้อความว่าง — ตกไป Tesseract")
        except TyphoonOCRError as exc:
            # ไม่โยนต่อ เพราะยังมีทางสำรอง ผู้ใช้ไม่ควรเห็น error ถ้ายังกู้ได้
            log.warning("Typhoon OCR ใช้ไม่ได้ (%s) — ตกไป Tesseract", exc)

    fallback = tesseract_pdf(data, tesseract_cmd)
    if fallback.strip():
        return DocumentText(
            _clean(fallback), "tesseract", True, n_pages,
            "อ่านด้วย Tesseract ความแม่นยำต่ำกว่า Typhoon OCR",
        )

    return DocumentText("", "none", False, n_pages, "อ่านข้อความจากไฟล์นี้ไม่ได้เลย")


def _read_image(
    data: bytes, *, typhoon: TyphoonOCR | None, tesseract_cmd: str
) -> DocumentText:
    if typhoon and typhoon.available:
        try:
            text = typhoon.ocr_image(data)
            if text.strip():
                return DocumentText(_clean(text), "typhoon_ocr", True, 1)
        except TyphoonOCRError as exc:
            log.warning("Typhoon OCR (รูปภาพ) ใช้ไม่ได้ (%s) — ตกไป Tesseract", exc)

    text = tesseract_image(data, tesseract_cmd)
    if text.strip():
        return DocumentText(_clean(text), "tesseract", True, 1)
    return DocumentText("", "none", False, 1, "อ่านข้อความจากรูปนี้ไม่ได้เลย")


def _read_docx(data: bytes) -> DocumentText:
    import io

    try:
        import docx  # python-docx
    except ImportError:
        return DocumentText("", "none", False, 0, "ยังไม่ได้ติดตั้ง python-docx")

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        return DocumentText("", "none", False, 0, f"เปิดไฟล์ .docx ไม่ได้ ({type(exc).__name__})")

    parts = [p.text for p in document.paragraphs]
    # ตารางใน .docx เก็บทักษะไว้บ่อย (ตารางโปรเจกต์ ตารางภาษา) ถ้าไม่ดึงจะหายไปเงียบ ๆ
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))

    return DocumentText(_clean("\n".join(parts)), "docx", False, 0)
